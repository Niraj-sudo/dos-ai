import streamlit as st
from transformers import AutoModelForCausalLM, AutoTokenizer,pipeline
from PyPDF2 import PdfReader
from langchain import  LLMChain, PromptTemplate
from langchain.llms import HuggingFacePipeline
import base64
from docx import Document
from huggingface_hub import login
import torch
from PIL import Image
import os
import secret
# Check if a GPU is available and set the deviceimport appropriately
login(secret.hugging_token)
#dev = torch.device("cuda" if torch.cuda.is_available() else "cpu")
def get_pdf_text(pdf_docs):
    text = ""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            text += page.extract_text()
        text=text.replace('\n', ' ')
    return text
def model():
    # Define Model ID
    max_memory_mapping = {0: "39GB", 1: "39GB", 2:"39GB"}
    model = AutoModelForCausalLM.from_pretrained("OpenAssistant/llama2-13b-orca-8k-3319",  device_map="auto",torch_dtype=torch.float16, max_memory=max_memory_mapping)
    # Load Tokenizer
    tokenizer = AutoTokenizer.from_pretrained("OpenAssistant/llama2-13b-orca-8k-3319", use_fast=False, stride=200)
    # Build HF Transformers pipeline 
    pipelines = pipeline(
        "text-generation", 
        model=model,
        tokenizer=tokenizer,
        device_map="auto",
        max_length=700,
        do_sample=True,
        top_k=10,
        num_return_sequences=1,
        eos_token_id=tokenizer.eos_token_id
    )
    return pipelines
def style_llm(texts, pipelines, style_options):
    
    template=f'''Act as a Editor and {style_options} style formatting. 1. Use active voice\n\t  2. Use a standard [subject] [verb] [object] sentence\n\t3. Don't Use personal pronouns\n\t4. Keep sentences and paragraphs concise\n\t5. Simplify complex information\n\t6. Use transition words\n\t7. Define uncommon terms and use headers and bulleted lists wherevr possible\nTarget Text for revision:\n
    '''
    prompt_template = f"{template} here is my draft proposal"+'{draft proposal}'
    llm = HuggingFacePipeline(pipeline=pipelines)
    llm_chain = LLMChain(llm=llm, prompt=PromptTemplate.from_template(prompt_template))
    new_texts=llm_chain(inputs={"draft proposal": f"{texts}"})
    return new_texts['text']
def get_binary_file_downloader_html(bin_file, file_label='File'):
    with open(bin_file, 'rb') as f:
        data = f.read()
    bin_str = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{bin_file}">Download {file_label}</a>'
    return href
image_url='https://th.bing.com/th/id/R.476c537795d84ee913dcb911e98a291c?rik=PXSzhkUUeNQ7BA&riu=http%3a%2f%2fwww2.deloitte.com%2fcontent%2fdam%2fDeloitte%2fin%2fImages%2fpromo_images%2fin-deloitte-logo-1x1-noexp.png&ehk=8vKWv4KVIHK2O9j%2faqQ8eu7pyxrogsOA36ILR0O8ZNU%3d&risl=&pid=ImgRaw&r=0&sres=1&sresct=1'
image_url2='/home/AD/npaneru/deploy_repo/dos-ai/logo3.png'
def chunk_string(string, chunk_size=400):
    return [string[i:i + chunk_size] for i in range(0, len(string), chunk_size)]
def main():
    st.set_page_config(page_title="Styles Guide Applicator",
                       page_icon=image_url)
    st.image(image_url2, width=110)
    st.header("Style Guide Applicator")
    if "pipelines" not in st.session_state:
        st.session_state.pipelines = None
    st.write('1. Select and upload your PDF(s)')
    pdf_docs = st.file_uploader(" ", accept_multiple_files=True)
    st.write('2. Select a style guide to apply to PDF(s)')
    style_options = ["Chicago Manual of Style", "APA Style", "MLA Style", "Other Style"]
    selected_style = st.selectbox("Select a style:", style_options)
    if selected_style == selected_style:
        if st.button('Submit'):
            with st.spinner("Processing"):
                text= get_pdf_text(pdf_docs)
                if len(text)<2:
                    st.write('The PDF does not contain any readable text. This might be because the PDF contains images of text rather than actual text content')
                    return 'Try Again'
                st.session_state.pipelines=model()
                if len(text)<450:
                    text=style_llm(texts=text,pipelines=st.session_state.pipelines, style_options=selected_style)
                else:
                    chunks = chunk_string(text)#split it into chucks
                    text=''
                    for chunk in chunks:
                        text=text+style_llm(texts=chunk,pipelines=st.session_state.pipelines, style_options=selected_style)
                doc = Document()
                style = doc.styles['Normal']
                font = style.font
                font.name = 'TimeRoman'
                doc.add_heading(f' {selected_style} Applicator', 0)
                doc.add_paragraph(text)
                doc.add_page_break()
                name=f'{selected_style} format Word Styling document'
                doc.save(f'{name}.docx')
                st.write('Click the link to download the document with Chicago style formating')
                st.markdown(get_binary_file_downloader_html(f'{name}.docx', f'{name}'), unsafe_allow_html=True)
                st.write('Here is a sample of AI translation:')
                st.write(text)
if __name__ == '__main__':
    main()
    
    
